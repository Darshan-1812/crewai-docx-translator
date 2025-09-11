import re
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def _apply_inline_formatting(paragraph, text):
	bold_pattern = re.compile(r"\*\*(.+?)\*\*")
	italic_pattern = re.compile(r"\*(.+?)\*")

	pos = 0
	segments = []
	for m in bold_pattern.finditer(text):
		start, end = m.span()
		if start > pos:
			segments.append((text[pos:start], None))
		segments.append((m.group(1), "bold"))
		pos = end
	if pos < len(text):
		segments.append((text[pos:], None))

	for seg_text, kind in segments:
		if kind == "bold":
			run = paragraph.add_run(seg_text)
			run.bold = True
		else:
			pos2 = 0
			for it in italic_pattern.finditer(seg_text):
				s2, e2 = it.span()
				if s2 > pos2:
					paragraph.add_run(seg_text[pos2:s2])
				run = paragraph.add_run(it.group(1))
				run.italic = True
				pos2 = e2
			if pos2 < len(seg_text):
				paragraph.add_run(seg_text[pos2:])


def markdown_to_docx(markdown_text: str, output_path: str) -> None:
	doc = Document()

	for raw_line in markdown_text.splitlines():
		line = raw_line.rstrip()
		if not line:
			doc.add_paragraph("")
			continue

		if line.startswith("### "):
			p = doc.add_paragraph()
			p.style = "Heading 3"
			_apply_inline_formatting(p, line[4:])
			continue
		if line.startswith("## "):
			p = doc.add_paragraph()
			p.style = "Heading 2"
			_apply_inline_formatting(p, line[3:])
			continue
		if line.startswith("# "):
			p = doc.add_paragraph()
			p.style = "Heading 1"
			_apply_inline_formatting(p, line[2:])
			continue

		if line.startswith("|") and line.endswith("|") and "|" in line[1:]:
			cells = [c.strip() for c in line.strip("|").split("|")]
			p = doc.add_paragraph()
			p.alignment = WD_ALIGN_PARAGRAPH.LEFT
			for idx, cell in enumerate(cells):
				r = p.add_run(cell)
				r.font.size = Pt(10)
				if idx < len(cells) - 1:
					p.add_run("\t")
			continue

		p = doc.add_paragraph()
		_apply_inline_formatting(p, line)

	doc.save(output_path)

