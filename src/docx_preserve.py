from __future__ import annotations
from dataclasses import dataclass
from typing import List, Tuple, Dict, Optional
from docx import Document
from docx.text.paragraph import Paragraph
from docx.table import _Cell, Table
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import logging


@dataclass
class TextUnit:
	path: Tuple[int, ...]
	text: str
	type: str  # 'paragraph' | 'table_cell' | 'footnote'
	footnote_id: Optional[int] = None  # For footnote references


@dataclass
class TokenStats:
	input_tokens: int = 0
	output_tokens: int = 0
	total_cost: float = 0.0


def extract_text_units(doc_path: str) -> Tuple[List[TextUnit], Document]:
	doc = Document(doc_path)
	units: List[TextUnit] = []

	# Extract paragraphs
	for p_idx, p in enumerate(doc.paragraphs):
		text = p.text or ""
		if text.strip():
			units.append(TextUnit(path=(p_idx,), text=text, type="paragraph"))

	# Extract table cells
	for t_idx, tbl in enumerate(doc.tables):
		for r_idx, row in enumerate(tbl.rows):
			for c_idx, cell in enumerate(row.cells):
				for p_idx, p in enumerate(cell.paragraphs):
					text = p.text or ""
					if text.strip():
						units.append(TextUnit(path=(t_idx, r_idx, c_idx, p_idx), text=text, type="table_cell"))

	# Extract footnotes
	footnote_units = extract_footnotes(doc)
	units.extend(footnote_units)

	return units, doc


def extract_footnotes(doc: Document) -> List[TextUnit]:
	"""Extract footnote content while preserving footnote IDs and references."""
	units = []
	
	try:
		# Access footnotes part
		if hasattr(doc.part, 'footnotes_part') and doc.part.footnotes_part:
			footnotes_part = doc.part.footnotes_part
			footnotes_xml = footnotes_part.element
			
			# Find all footnote elements
			for footnote_elem in footnotes_xml.findall('.//w:footnote', {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}):
				footnote_id = footnote_elem.get(qn('w:id'))
				if footnote_id:
					footnote_id = int(footnote_id)
					
					# Extract text from footnote paragraphs
					for p_idx, p_elem in enumerate(footnote_elem.findall('.//w:p', {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})):
						text_parts = []
						for t_elem in p_elem.findall('.//w:t', {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}):
							if t_elem.text:
								text_parts.append(t_elem.text)
						
						footnote_text = ''.join(text_parts).strip()
						if footnote_text:
							units.append(TextUnit(
								path=('footnote', footnote_id, p_idx),
								text=footnote_text,
								type="footnote",
								footnote_id=footnote_id
							))
	except Exception as e:
		logging.warning(f"Could not extract footnotes: {e}")
	
	return units


def replace_text_in_document(doc: Document, original_to_translated: List[Tuple[TextUnit, str]]) -> Document:
	for unit, translated in original_to_translated:
		if unit.type == "paragraph" and len(unit.path) == 1:
			p_idx = unit.path[0]
			p: Paragraph = doc.paragraphs[p_idx]
			for _ in range(len(p.runs)):
				p.runs[0]._r.getparent().remove(p.runs[0]._r)
			p.add_run(translated)
		elif unit.type == "table_cell" and len(unit.path) == 4:
			t_idx, r_idx, c_idx, p_idx = unit.path
			tbl: Table = doc.tables[t_idx]
			cell: _Cell = tbl.rows[r_idx].cells[c_idx]
			p: Paragraph = cell.paragraphs[p_idx]
			for _ in range(len(p.runs)):
				p.runs[0]._r.getparent().remove(p.runs[0]._r)
			p.add_run(translated)
		elif unit.type == "footnote" and len(unit.path) == 3:
			replace_footnote_text(doc, unit, translated)
	return doc


def replace_footnote_text(doc: Document, unit: TextUnit, translated: str):
	"""Replace footnote text while preserving footnote structure and references."""
	try:
		if hasattr(doc.part, 'footnotes_part') and doc.part.footnotes_part:
			footnotes_part = doc.part.footnotes_part
			footnotes_xml = footnotes_part.element
			
			# Find the specific footnote by ID
			for footnote_elem in footnotes_xml.findall('.//w:footnote', {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}):
				footnote_id = footnote_elem.get(qn('w:id'))
				if footnote_id and int(footnote_id) == unit.footnote_id:
					# Find the specific paragraph within the footnote
					p_elements = footnote_elem.findall('.//w:p', {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
					p_idx = unit.path[2]  # The paragraph index within the footnote
					
					if p_idx < len(p_elements):
						p_elem = p_elements[p_idx]
						
						# Remove existing text runs
						for t_elem in p_elem.findall('.//w:t', {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}):
							t_elem.text = ""
						
						# Add new text
						if p_elem.find('.//w:r', {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}) is not None:
							# Use existing run
							run_elem = p_elem.find('.//w:r', {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
							t_elem = run_elem.find('.//w:t', {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
							if t_elem is not None:
								t_elem.text = translated
						else:
							# Create new run
							run_elem = OxmlElement('w:r')
							t_elem = OxmlElement('w:t')
							t_elem.text = translated
							run_elem.append(t_elem)
							p_elem.append(run_elem)
					break
	except Exception as e:
		logging.warning(f"Could not replace footnote text: {e}")


def estimate_tokens(text: str) -> int:
	"""Rough token estimation: ~4 characters per token for most languages."""
	return len(text) // 4


def calculate_optimal_batch_size(units: List[TextUnit], max_tokens_per_batch: int = 2000) -> int:
	"""Calculate optimal batch size based on token limits."""
	if not units:
		return 1
	
	total_tokens = sum(estimate_tokens(unit.text) for unit in units)
	if total_tokens <= max_tokens_per_batch:
		return len(units)
	
	# Start with a conservative batch size and adjust
	batch_size = max(1, len(units) // (total_tokens // max_tokens_per_batch + 1))
	
	# Verify the batch size doesn't exceed token limits
	while batch_size > 1:
		batch_tokens = sum(estimate_tokens(unit.text) for unit in units[:batch_size])
		if batch_tokens <= max_tokens_per_batch:
			break
		batch_size = max(1, batch_size - 1)
	
	return batch_size


def log_token_stats(stats: TokenStats, batch_num: int, total_batches: int):
	"""Log detailed token statistics."""
	logging.info(f"Batch {batch_num}/{total_batches} - Input: {stats.input_tokens} tokens, "
				f"Output: {stats.output_tokens} tokens, Cost: ${stats.total_cost:.4f}")

